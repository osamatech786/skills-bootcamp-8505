import streamlit as st
from datetime import datetime, date
from PIL import Image as PILImage
import numpy as np
# import io
from docx import Document
from docx.shared import Inches
from streamlit_drawable_canvas import st_canvas
import smtplib
from email.message import EmailMessage
import shutil
import re
import time
from dotenv import load_dotenv
import os

# Set page configuration with a favicon
st.set_page_config(
    page_title="Surrey H&S Bootcamp",
    page_icon="https://lirp.cdn-website.com/d8120025/dms3rep/multi/opt/social-image-88w.png", 
    layout="centered"  # "centered" or "wide"
)

# add render support along with st.secret
def get_secret(key):
    try:
        load_dotenv()
        # Attempt to get the secret from environment variables
        secret = os.environ.get(key)
        if secret is None:
            raise ValueError("Secret not found in environment variables")
        return secret
    except (ValueError, TypeError) as e:
        # If an error occurs, fall back to Streamlit secrets
        if hasattr(st, 'secrets'):
            return st.secrets.get(key)
        # If still not found, return None or handle as needed
        return None

# Initialize session state
if 'step' not in st.session_state: st.session_state.step = 1
if 'submission_done' not in st.session_state: st.session_state.submission_done = False
if 'selected_option' not in st.session_state: st.session_state.selected_option = "    "  # Default value for the dropdown
if 'hear_about' not in st.session_state: st.session_state.hear_about = "Self-referral"  # Default value for the dropdown
if 'hother_source' not in st.session_state: st.session_state.hother_source = ''  # Default for the "Other" text input
# Initialize session state for Step 2
if 'title' not in st.session_state: st.session_state.title = "Mr"  # Default title
if 'sir_name' not in st.session_state: st.session_state.sir_name = ''  # Default surname
if 'first_name' not in st.session_state: st.session_state.first_name = ''  # Default first name
if 'preferred_name' not in st.session_state: st.session_state.preferred_name = ''  # Default preferred name
if 'previous_name' not in st.session_state: st.session_state.previous_name = ''  # Default previous name
if 'home_address' not in st.session_state: st.session_state.home_address = ''  # Default home address
if 'postcode' not in st.session_state: st.session_state.postcode = ''  # Default postcode
if 'previous_postcode_country' not in st.session_state: st.session_state.previous_postcode_country = ''  # Default previous postcode or country
if 'dob' not in st.session_state: st.session_state.dob = None
if 'ni_number' not in st.session_state: st.session_state.ni_number = ''  # Default National Insurance number
if 'gender' not in st.session_state: st.session_state.gender = "Male"  # Default gender
if 'home_number' not in st.session_state: st.session_state.home_number = ''  # Default home phone number
if 'mobile_number' not in st.session_state: st.session_state.mobile_number = ''  # Default mobile number
if 'email' not in st.session_state: st.session_state.email = ''  # Default email
# Initialize session state for Step 3
ethnicity_options = {
    'White': {
        'English/ Welsh/ Scottish/ N Irish/ British': '31',
        'Irish': '32',
        'Roma, Gypsy or Irish Traveller': '33',
        'Any other white background': '34'
    },
    'Mixed/ Multiple ethnic group': {
        'White and Black Caribbean': '35',
        'White and Black African': '36',
        'White and Asian': '37',
        'Any other mixed/ multiple ethnic background': '38'
    },
    'Asian/ Asian British': {
        'Bangladeshi': '41',
        'Chinese': '42',
        'Indian': '39',
        'Pakistani': '40',
        'Any other Asian background': '43'
    },
    'Black/ African/ Caribbean/ Black British': {
        'African': '44',
        'Caribbean': '45',
        'Any Other Black/ African/ Caribbean background': '46'
    },
    'Other Ethnic Group': {
        'Arab': '47',
        'Any other ethnic group': '48'
    }
}
if 'ethnicity_category' not in st.session_state: st.session_state.ethnicity_category = list(ethnicity_options.keys())[0]  # Default to first category
# Initialize ethnicity only if ethnicity_category is already defined
if 'ethnicity' not in st.session_state: st.session_state.ethnicity = list(ethnicity_options[st.session_state.ethnicity_category].keys())[0]  # Default to first ethnicity in the default category
if 'ethnicity_code' not in st.session_state: st.session_state.ethnicity_code = 31  # Default to first code
if 'ethnicity_vars' not in st.session_state: st.session_state.ethnicity_vars = {f'ethnicity_{i}': '' for i in range(31, 49)}  # Initialize all ethnicity variables
if 'ph59' not in st.session_state: st.session_state.ph59 = ''  # Default for criminal conviction
if 'ph60' not in st.session_state: st.session_state.ph60 = ''  # Default for criminal conviction
if 'criminal_conviction' not in st.session_state: st.session_state.criminal_conviction = "No"  # Default for criminal conviction radio
if 'ph61' not in st.session_state: st.session_state.ph61 = ''  # Default for caring for children
if 'ph62' not in st.session_state: st.session_state.ph62 = ''  # Default for caring for children
if 'caring_children' not in st.session_state: st.session_state.caring_children = "No"  # Default for caring for children radio
# Initialize session state for Step 4
if 'emergency_contact_name' not in st.session_state: st.session_state.emergency_contact_name = ''  # Default for emergency contact name
if 'emergency_contact_relationship' not in st.session_state: st.session_state.emergency_contact_relationship = ''  # Default for emergency contact relationship
if 'emergency_contact_phone' not in st.session_state: st.session_state.emergency_contact_phone = ''  # Default for emergency contact mobile number
if 'home_tel_no' not in st.session_state: st.session_state.home_tel_no = ''  # Default for emergency contact home telephone number
# Initialize session state for Step 5
if 'ph63' not in st.session_state: st.session_state.ph63 = ''  # Entry Level
if 'ph64' not in st.session_state: st.session_state.ph64 = ''  # Qualifications below Level 1
if 'ph65' not in st.session_state: st.session_state.ph65 = ''  # Level 1
if 'ph66' not in st.session_state: st.session_state.ph66 = ''  # Full Level 2
if 'ph67' not in st.session_state: st.session_state.ph67 = ''  # Full Level 3
if 'ph68' not in st.session_state: st.session_state.ph68 = ''  # Level 4
if 'ph69' not in st.session_state: st.session_state.ph69 = ''  # Level 5
if 'ph70' not in st.session_state: st.session_state.ph70 = ''  # Level 6
if 'ph71' not in st.session_state: st.session_state.ph71 = ''  # Level 7 or above
if 'ph72' not in st.session_state: st.session_state.ph72 = ''  # Other qualification
if 'ph73' not in st.session_state: st.session_state.ph73 = ''  # No qualifications
if 'ph74' not in st.session_state: st.session_state.ph74 = ''  # Not known
# Initialize session state for Step 6
for i in range(75, 84):
    if f'ph{i}' not in st.session_state: st.session_state[f'ph{i}'] = ''  # Initialize placeholders for employment status options
# Additional placeholders for employment details
if 'ph79a' not in st.session_state: st.session_state.ph79a = ''  # 0 – 10 Hours (Self-employed)
if 'ph79b' not in st.session_state: st.session_state.ph79b = ''  # 11 – 20 Hours (Self-employed)
if 'ph79c' not in st.session_state: st.session_state.ph79c = ''  # 21 – 30 Hours (Self-employed)
if 'ph79d' not in st.session_state: st.session_state.ph79d = ''  # 31+ Hours (Self-employed)
# Initialize placeholders for other options
for i in range(83, 88):
    if f'ph{i}' not in st.session_state: st.session_state[f'ph{i}'] = ''  # Unemployment duration options
for i in range(88, 93):
    if f'ph{i}' not in st.session_state: st.session_state[f'ph{i}'] = ''  # Benefit options
for i in range(93, 96):
    if f'ph{i}' not in st.session_state: st.session_state[f'ph{i}'] = ''  # Employer details
if 'ph96' not in st.session_state: st.session_state.ph96 = 0.0
if 'ph97y' not in st.session_state: st.session_state.ph97y = ''  # Attending Bootcamp via Employer (Yes/No)
if 'ph97n' not in st.session_state: st.session_state.ph97n = ''  # Attending Bootcamp via Employer (Yes/No)
for i in range(98, 102):
    if f'ph{i}' not in st.session_state: st.session_state[f'ph{i}'] = ''  # Work alongside bootcamp options
for i in range(102, 112):
    if f'ph{i}' not in st.session_state: st.session_state[f'ph{i}'] = ''  # Most recent occupation options
for i in range(112, 121):
    if f'ph{i}' not in st.session_state: st.session_state[f'ph{i}'] = ''  # Industry/sector options
if 'ph120a' not in st.session_state: st.session_state.ph120a = ''  # Other services (Specify)
# Initialize session state for Step 7
for i in range(121, 147):
    if f'ph{i}' not in st.session_state: st.session_state[f'ph{i}'] = ''  # Initialize placeholders for disabilities, learning difficulties, and health problems
if 'ph123a' not in st.session_state: st.session_state.ph123a = ''  # Other Specify
if 'impactful_condition' not in st.session_state: st.session_state.impactful_condition = ''  # Default for most impactful condition
if 'confidential_interview' not in st.session_state: st.session_state.confidential_interview = ''  # Default for confidential interview request
# Initialize session state for Step 8
for i in range(147, 154):
    if f'ph{i}' not in st.session_state: st.session_state[f'ph{i}'] = ''  # Initialize placeholders for contact and marketing information
if 'other_source' not in st.session_state: st.session_state.other_source = ''  # Default for other source specification
# Initialize session state for Step 9
for i in range(154, 160):
    if f'ph{i}' not in st.session_state: st.session_state[f'ph{i}'] = ''  # Initialize placeholders for learner declaration and commitment



def last():
    st.session_state.clear()

def is_valid_email(email):
    # Comprehensive regex for email validation
    pattern = r'''
        ^                         # Start of string
        (?!.*[._%+-]{2})          # No consecutive special characters
        [a-zA-Z0-9._%+-]{1,64}    # Local part: allowed characters and length limit
        (?<![._%+-])              # No special characters at the end of local part
        @                         # "@" symbol
        [a-zA-Z0-9.-]+            # Domain part: allowed characters
        (?<![.-])                 # No special characters at the end of domain
        \.[a-zA-Z]{2,}$           # Top-level domain with minimum 2 characters
    '''
    
    # Match the entire email against the pattern
    return re.match(pattern, email, re.VERBOSE) is not None

# Function to send email with attachments (Handle Local + Uploaded)
def send_email_with_attachments(sender_email, sender_password, receiver_email, subject, body, files=None, local_file_path=None):
    msg = EmailMessage()
    msg['From'] = sender_email
    msg['To'] = ", ".join(receiver_email)
    msg['Subject'] = subject
    msg.set_content(body, subtype='html')

    # Attach uploaded files
    if files:
        for uploaded_file in files:
            uploaded_file.seek(0)  # Move to the beginning of the UploadedFile
            msg.add_attachment(uploaded_file.read(), maintype='application', subtype='octet-stream', filename=uploaded_file.name)

    # Attach local file if specified
    if local_file_path:
        with open(local_file_path, 'rb') as f:
            file_data = f.read()
            file_name = local_file_path.split('/')[-1]
            msg.add_attachment(file_data, maintype='application', subtype='octet-stream', filename=file_name)

    # Use the SMTP server for sending the email
    with smtplib.SMTP('smtp.office365.com', 587) as server:
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)

def is_signature_drawn(signature):
    # Check if signature is None or an empty numpy array
    if signature is None:
        return False
    # Ensure it is a numpy array and has content
    if isinstance(signature, np.ndarray) and signature.size > 0:
        # Additional check: if the array is not just empty white pixels
        # Assuming white background is [255, 255, 255] in RGB
        if np.all(signature == 255):
            return False
        return True
    return False

def resize_image_to_fit_cell(image, max_width, max_height):
    width, height = image.size
    aspect_ratio = width / height

    if width > max_width:
        width = max_width
        height = int(width / aspect_ratio)

    if height > max_height:
        height = max_height
        width = int(height * aspect_ratio)

    return image.resize((width, height))


def replace_placeholders(template_file, modified_file, placeholder_values, signature_path):
    try:
        print(f"Copying template file '{template_file}' to '{modified_file}'...")
        shutil.copy(template_file, modified_file)

        print(f"Opening document '{modified_file}'...")
        doc = Document(modified_file)

        # Function to convert value to string, handling datetime.date objects
        def convert_to_str(value):
            if isinstance(value, date):
                return value.strftime('%Y-%m-%d')  # Convert date to string
            return str(value)  # Convert other types to string

        # Compile regular expressions for all placeholders
        placeholders = {re.escape(key): convert_to_str(value) for key, value in placeholder_values.items()}
        placeholders_pattern = re.compile(r'\b(' + '|'.join(placeholders.keys()) + r')\b')

        # Replace placeholders in paragraphs
        print("Replacing placeholders in paragraphs...")
        for para in doc.paragraphs:
            original_text = para.text
            updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], para.text)
            if original_text != updated_text:
                print(f"Updated paragraph text: '{original_text}' -> '{updated_text}'")
                para.text = updated_text

        # Replace placeholders in tables
        print("Replacing placeholders in tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], para.text)
                        if original_text != updated_text:
                            print(f"Updated table cell text: '{original_text}' -> '{updated_text}'")
                            para.text = updated_text

                    # Inspect cell runs
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run_text = run.text
                            run_updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], run_text)
                            if run_text != run_updated_text:
                                print(f"Updated run text in table cell: '{run_text}' -> '{run_updated_text}'")
                                run.text = run_updated_text

        # Check and handle signature placeholder
        print("Inspecting document for 'ph_signature' placeholder...")
        signature_placeholder_found = False

        # Check paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()  # Remove any extra spaces around text
            while 'ph_signature' in para_text:
                print(f"Found 'ph_signature' in paragraph: '{para_text}'")
                para_text = para_text.replace('ph_signature', '').strip()  # Remove 'ph_signature' and any leading/trailing spaces
                para.text = para_text
                resized_image_path = 'resized_signature_image.png'
                
                try:
                    # Open and resize the image
                    print(f"Opening image file: {signature_path}")
                    resized_image = PILImage.open(signature_path)
                    print(f"Original image size: {resized_image.size}")
                    resized_image = resize_image_to_fit_cell(resized_image, 200, 50)
                    resized_image.save(resized_image_path)  # Save resized image to a file
                    print(f"Resized image saved to: {resized_image_path}")
                    
                    # Add picture to the paragraph
                    print(f"Adding picture to paragraph from path: {resized_image_path}")
                    para.add_run().add_picture(resized_image_path, width=Inches(2))
                    print("Inserted signature image into paragraph.")
                    signature_placeholder_found = True
                except Exception as img_e:
                    print(f"An error occurred with image processing: {img_e}")

        # Check table cells again in case the placeholder was missed
        if not signature_placeholder_found:
            print("Checking table cells for 'ph_signature'...")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            para_text = para.text.strip()
                            while 'ph_signature' in para_text:
                                print(f"Found 'ph_signature' in table cell paragraph: '{para_text}'")
                                para_text = para_text.replace('ph_signature', '').strip()
                                para.text = para_text
                                resized_image_path = 'resized_signature_image.png'
                                
                                try:
                                    # Open and resize the image
                                    print(f"Opening image file: {signature_path}")
                                    resized_image = PILImage.open(signature_path)
                                    print(f"Original image size: {resized_image.size}")
                                    resized_image = resize_image_to_fit_cell(resized_image, 200, 50)
                                    resized_image.save(resized_image_path)  # Save resized image to a file
                                    print(f"Resized image saved to: {resized_image_path}")
                                    
                                    # Add picture to the table cell
                                    print(f"Adding picture to table cell from path: {resized_image_path}")
                                    para.add_run().add_picture(resized_image_path, width=Inches(2))
                                    print("Inserted signature image into table cell.")
                                    signature_placeholder_found = True
                                except Exception as img_e:
                                    print(f"An error occurred with image processing: {img_e}")

        if not signature_placeholder_found:
            print("No signature placeholder found.")

        # Save the modified document
        print(f"Saving modified document '{modified_file}'...")
        doc.save(modified_file)
        print(f"Document modification complete: '{modified_file}'")

    except Exception as e:
        print(f"An error occurred: {e}")

    # file download button
    with open(modified_file, 'rb') as f:
        file_contents = f.read()
        st.download_button(
            label="Download Your Response",
            data=file_contents,
            file_name=modified_file,
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

def calculate_age(born):
    today = date.today()
    return today.year - born.year - ((today.month, today.day) < (born.month, born.day))

if 'files' not in st.session_state:
    st.session_state.files = []

# Define a function to calculate progress and percentage
def get_progress(step, total_steps=14):
    return int((step / total_steps) * 100)




# Define the total number of steps
total_steps = 10
# Calculate the current progress
progress = get_progress(st.session_state.step, total_steps)
# Display the progress bar and percentage
st.write(f"Progress: {progress}%")
st.progress(progress)




# Define different steps
if st.session_state.step == 1:
    st.image('resources/header-wihout-bg.png', use_column_width=True)

    st.title("Skills Bootcamp Enrollment and Registration Document")
    st.write("Provider: Prevista Ltd. | Sponsor: Surrey County Council | Website: www.prevista.co.uk")
    st.write("________________________________________")
    st.write("**Application and Enrollment Form**")

    # Add question with a dropdown menu
    support_options = [
        "    ", 
        "Self Completing",
        "Innovator Recruitment Team",
        "Catalyst Recruitment Team",
        "Guildford JCP",
        "Camberley JCP",
        "Epsom JCP",
        "Woking JCP",
        "Redhill JCP",
        "Staines JCP",
        "Leatherhead JCP",
        "Croydon JCP",
        "Surrey County Council",
        "Surrey Employment and Skills Board (SESB)",
        "Federation of Small Businesses (FSB) Surrey",
        "Surrey Chambers of Commerce",
        "Voluntary Action South West Surrey",
        "Guildford Borough Council",
        "Woking Borough Council",
        "Surrey Choices",
        "Elmbridge Community Job Club",
        "Mole Valley Employment Group",
        "Surrey Lifelong Learning Partnership (SLLP)",
    ]

    st.session_state.selected_option = st.selectbox(
        "Who is supporting you to fill this form?", 
        support_options,
        index=support_options.index(st.session_state.selected_option)  # Set default value
    )

    hear_about_options = [
        "Self-referral", 
        "Jobcentre Plus (JCP)",
        "Local Council",
        "Online",
        "Word of Mouth",
        "Community Organization",
        "Employer or Training Provider",
        "Promotional Materials",
        "Other (please specify)"
    ]

    st.session_state.hear_about = st.selectbox(
        "Hear about this opportunity:", 
        hear_about_options,
        index=hear_about_options.index(st.session_state.hear_about)  # Set default value
    )

    # If the user selects "Other (please specify)", display an input field
    if st.session_state.hear_about == "Other (please specify)":
        st.session_state.hother_source = st.text_input("Please specify:", value=st.session_state.hother_source)  # Set default value
    else:
        st.session_state.hother_source = ''  # Reset if not selected

    st.write("""
    Please complete the upcoming sections to finalize your enrollment.
    """)

    if st.button("Next"):
        if st.session_state.selected_option != '    ':
            st.session_state.step = 2
            st.experimental_rerun()
        else:
            st.warning("Please Choose Valid Support Option.")


elif st.session_state.step == 2:
    st.title("> 1: Learner Information")
    
    st.session_state.title = st.radio(
        "Title",
        ["Mr", "Mrs", "Miss", "Ms"],
        index=["Mr", "Mrs", "Miss", "Ms"].index(st.session_state.title)  # Set default value
    )
    
    st.session_state.sir_name = st.text_input("Surname/Family Name", value=st.session_state.sir_name)
    st.session_state.first_name = st.text_input("First Name(s) in full", value=st.session_state.first_name)
    st.session_state.preferred_name = st.text_input("Preferred Name", value=st.session_state.preferred_name)
    st.session_state.previous_name = st.text_input("Previous Name (if applicable)", value=st.session_state.previous_name)
    st.session_state.home_address = st.text_input("Home Address", value=st.session_state.home_address)
    st.session_state.postcode = st.text_input("Home Postcode", value=st.session_state.postcode)
    st.session_state.previous_postcode_country = st.text_input("If you have changed address within the last 3 years, please provide previous UK Postcode / Country (if not living in the UK)", value=st.session_state.previous_postcode_country)
    
    # Check if dob is a string and convert it back to a date object
    if isinstance(st.session_state.get("dob"), str):
        st.session_state.dob = datetime.strptime(st.session_state.get("dob"), "%d-%m-%Y").date()

    # Date of Birth
    st.session_state.dob = st.date_input(
        label="Date of Birth",  # Label for the field
        value=st.session_state.get("dob"),  # Correctly access dob from session state
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date.today(),  # Maximum selectable date
        help="Choose a date",  # Tooltip text
        format='DD/MM/YYYY'
    )

    st.session_state.ni_number = st.text_input("National Insurance Number", value=st.session_state.ni_number)

    st.session_state.ph35m, st.session_state.ph35f = '', ''
    st.session_state.gender = st.radio(
        "Legal Sex as stated on passport/birth certificate:", 
        ["Male", "Female"],
        index=["Male", "Female"].index(st.session_state.gender)  # Set default value
    )

    # Update the session state based on the gender selection
    if st.session_state.gender == "Male":
        st.session_state.ph35m = 'x'
        st.session_state.ph35f = ''
    elif st.session_state.gender == "Female":
        st.session_state.ph35m = ''
        st.session_state.ph35f = 'x'
        
    st.session_state.home_number = st.text_input("Home Tel No", value=st.session_state.home_number)
    st.session_state.mobile_number = st.text_input("Mobile Number", value=st.session_state.mobile_number)
    st.session_state.email = st.text_input("Email Address", value=st.session_state.email)

    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        if is_valid_email(st.session_state.email):
            if (st.session_state.title and
                st.session_state.sir_name and
                st.session_state.first_name and
                st.session_state.preferred_name and
                st.session_state.home_address and
                st.session_state.postcode and
                st.session_state.dob and
                st.session_state.ni_number and
                st.session_state.home_number and
                st.session_state.mobile_number):

                # Convert the selected date to the desired string format (DD-MM-YYYY) only when proceeding to the next step
                st.session_state.dob = st.session_state.dob.strftime("%d-%m-%Y")

                # Check if dob is a string and convert it back to a date object
                if isinstance(st.session_state.get("dob"), str):
                    st.session_state.dob = datetime.strptime(st.session_state.get("dob"), "%d-%m-%Y").date()
                st.session_state.current_age = calculate_age(st.session_state.dob)
                st.session_state.current_age_text = 'Current Age: ' + str(st.session_state.current_age)
                st.text(st.session_state.current_age_text)
                
                st.session_state.step = 3
                st.experimental_rerun()
            else:
                st.warning("Please fill in all fields before proceeding.")    
        else:
            st.warning("Please enter a valid email address.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 1  # Go back to the previous step (Section 1)
        st.experimental_rerun()


elif st.session_state.step == 3:
    st.title("> 2: Please indicate your ethnic group")

    # Initialize ethnicity_vars if it doesn't exist
    if 'ethnicity_vars' not in st.session_state:
        st.session_state.ethnicity_vars = {f'ethnicity_{i}': '' for i in range(31, 49)}

    # Select ethnicity category
    st.session_state.ethnicity_category = st.selectbox(
        'Select Ethnicity Category', 
        list(ethnicity_options.keys()),
        index=list(ethnicity_options.keys()).index(st.session_state.ethnicity_category)  # Set default value
    )

    # Update the ethnicity selection based on the selected category
    if 'ethnicity' not in st.session_state or st.session_state.ethnicity not in ethnicity_options[st.session_state.ethnicity_category]:
        # Reset ethnicity to the first option of the new category if it's not valid
        st.session_state.ethnicity = list(ethnicity_options[st.session_state.ethnicity_category].keys())[0]

    # Select ethnicity based on updated category
    st.session_state.ethnicity = st.selectbox(
        'Select Ethnicity', 
        list(ethnicity_options[st.session_state.ethnicity_category].keys()),
        index=list(ethnicity_options[st.session_state.ethnicity_category].keys()).index(st.session_state.ethnicity)  # Set default value
    )

    # Retrieve and convert ethnicity code to integer
    ethnicity_code_str = ethnicity_options[st.session_state.ethnicity_category][st.session_state.ethnicity]
    st.session_state.ethnicity_code = int(ethnicity_code_str)  # Ensure it is an integer

    # Clear previous ethnicity selections in ethnicity_vars
    for key in range(31, 49):  # Adjust this range according to your ethnicity codes
        st.session_state.ethnicity_vars[f'ethnicity_{key}'] = ''  # Reset all options to empty string

    # Set the corresponding ethnicity variable to 'X'
    if st.session_state.ethnicity_code in range(31, 49):
        st.session_state.ethnicity_vars[f'ethnicity_{st.session_state.ethnicity_code}'] = 'X'

    # Reset previous selections for criminal conviction
    st.session_state.ph59, st.session_state.ph60 = '', ''
    if 'criminal_conviction' not in st.session_state:
        st.session_state.criminal_conviction = "No"  # Default value

    # Radio button for criminal conviction
    st.session_state.criminal_conviction = st.radio(
        "Do you have a criminal conviction (excluding minor motoring offences)?", 
        ["No", "Yes"], 
        index=["No", "Yes"].index(st.session_state.criminal_conviction)
    )

    # Update the session state based on the criminal_conviction selection
    if st.session_state.criminal_conviction == "Yes":
        st.session_state.ph59 = 'x'
        st.session_state.ph60 = ''
    elif st.session_state.criminal_conviction == "No":
        st.session_state.ph59 = ''
        st.session_state.ph60 = 'x'

    # Reset previous selections for caring for children
    st.session_state.ph61, st.session_state.ph62 = '', ''
    if 'caring_children' not in st.session_state:
        st.session_state.caring_children = "No"  # Default value

    # Radio button for caring for children or other adults
    st.session_state.caring_children = st.radio(
        "Are you currently caring for children or other adults?", 
        ["No", "Yes"], 
        index=["No", "Yes"].index(st.session_state.caring_children)
    )

    # Update the session state based on the caring_children selection
    if st.session_state.caring_children == "Yes":
        st.session_state.ph61 = 'x'
        st.session_state.ph62 = ''
    elif st.session_state.caring_children == "No":
        st.session_state.ph61 = ''
        st.session_state.ph62 = 'x'    

    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        # Proceed to the next step
        st.session_state.step = 4
        st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 2  # Go back to the previous step (Section 1)
        st.experimental_rerun()



elif st.session_state.step == 4:
    st.title("> 3: Emergency Contact Details")

    st.session_state.emergency_contact_name = st.text_input("Emergency Contact Name", value=st.session_state.emergency_contact_name)  # Set default value
    st.session_state.emergency_contact_relationship = st.text_input("Emergency Contact Relationship", value=st.session_state.emergency_contact_relationship)  # Set default value
    st.session_state.emergency_contact_phone = st.text_input("Emergency Contact Mobile Number", value=st.session_state.emergency_contact_phone)  # Set default value
    st.session_state.home_tel_no = st.text_input("Emergency Contact Home Tel No", value=st.session_state.home_tel_no)  # Set default value

    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        if (st.session_state.emergency_contact_name and
            st.session_state.emergency_contact_relationship and
            st.session_state.emergency_contact_phone and
            st.session_state.home_tel_no):
            st.session_state.step = 5
            st.experimental_rerun()
        else:
            st.warning("Please fill in all fields before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 3  # Go back to the previous step (Section 2)
        st.experimental_rerun()


elif st.session_state.step == 5:
    st.title("> 4: Prior Attainment/Highest Previous Qualifications")

    # Initialize placeholders if not already done
    if 'qualification_level' not in st.session_state:
        st.session_state.qualification_level = None  # Default to None or some valid default

    # Define radio button options
    options = {
        "Entry Level (Basic Entry Level, E)": "ph63",
        "Qualifications below Level 1 (Pre-entry)": "ph64",
        "Level 1 (5GCSEs D-G/3-1; 1 AS Level; GNVQ Foundation; BTEC First Certificate)": "ph65",
        "Full Level 2 (5 GCSEs A*-C/9-4; NVQ2; 2 or 3 AS Levels; GNVQ Intermediate; BTEC First Diploma)": "ph66",
        "Full Level 3 (4 AS Level; 2 A2/A Level; NVQ3; BTEC Diploma/Extended Diploma/Access to HE)": "ph67",
        "Level 4 (Certificate of Higher Education; HNC)": "ph68",
        "Level 5 (Foundation Degree; HND)": "ph69",
        "Level 6 (Bachelor’s Degree; Graduate qualification)": "ph70",
        "Level 7 or above (Master’s Degree; Postgraduate qualification; Doctorate)": "ph71",
        "Other qualification: level not known": "ph72",
        "No qualifications": "ph73",
        "Not known": "ph74"
    }

    # Retrieve the previously selected option if available
    selected_option = None
    for option, placeholder in options.items():
        if st.session_state.get(placeholder) == 'X':
            selected_option = option
            break

    # Create a radio button and store 'X' in the selected placeholder
    selected_option = st.radio("Select your qualification level:", list(options.keys()), index=list(options.keys()).index(selected_option) if selected_option else 0)

    # Update the corresponding placeholder with 'X'
    if selected_option:
        # Reset all placeholders before setting the current one
        for key in options.values():
            st.session_state[key] = ''  # Reset all qualifications to an empty string
        st.session_state[options[selected_option]] = 'X'  # Set selected option to 'X'
        st.session_state.qualification_level = selected_option  # Store the selected option

    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        st.session_state.step = 6
        st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 4  # Go back to the previous step (Section 3)
        st.experimental_rerun()



elif st.session_state.step == 6:
    st.title("> 5: Employment Information")

    # Initialize placeholders for employment status options
    st.session_state.update({f'ph{i}': st.session_state.get(f'ph{i}', '') for i in range(75, 122)})  # Make sure all relevant placeholders are initialized

    # Define radio button options and corresponding placeholders
    employment_options = {
        "0 – 10 Hours (Paid employment)": "ph75",
        "11 – 20 Hours (Paid employment)": "ph76",
        "21 – 30 Hours (Paid employment)": "ph77",
        "31+ Hours (Paid employment)": "ph78",
        "0 – 10 Hours (Self-employed)": "ph79a",
        "11 – 20 Hours (Self-employed)": "ph79b",
        "21 – 30 Hours (Self-employed)": "ph79c",
        "31+ Hours (Self-employed)": "ph79d",
        "Not in paid employment & looking for work": "ph80",
        "Not in paid employment & not looking for work": "ph81",
        "In full-time education or training prior to enrolment": "ph82"
    }

    # Retrieve the previously selected employment status if available
    selected_employment = None
    for option, placeholder in employment_options.items():
        if st.session_state.get(placeholder) == 'X':
            selected_employment = option
            break

    # Display radio buttons and set default value
    selected_employment = st.radio(
        "On the day prior to this course, what is your employment status? (Please tick ONE box)",
        list(employment_options.keys()),
        index=list(employment_options.keys()).index(selected_employment) if selected_employment else 0
    )

    # Update the corresponding placeholder with 'X'
    if selected_employment:
        # Reset all placeholders before setting the current one
        for key in employment_options.values():
            st.session_state[key] = ''  # Reset all employment options to empty string
        st.session_state[employment_options[selected_employment]] = 'X'  # Set selected option to 'X'

    # Conditionally display additional sections if "Paid employment" | or "not in paid employment" is selected
    paid_employment_options = [
        "0 – 10 Hours (Paid employment)",
        "11 – 20 Hours (Paid employment)",
        "21 – 30 Hours (Paid employment)",
        "31+ Hours (Paid employment)"
    ]

    unpaid_employment_options = [
        "Not in paid employment & looking for work",
        "Not in paid employment & not looking for work",
        "In full-time education or training prior to enrolment",
    ]

    if selected_employment in paid_employment_options:
        st.subheader("Additional Information for Paid Employment")

        # Additional fields for Paid Employment
        name_of_employer = st.text_input("Name of Employer", value=st.session_state.get('ph93', ''))
        employer_postcode = st.text_input("Postcode", value=st.session_state.get('ph94', ''))
        current_job_role = st.text_input("Current Job Role", value=st.session_state.get('ph95', ''))
        current_hourly_rate = st.number_input(
            "Current Hourly Rate", 
            min_value=0.0, 
            format="%.2f", 
            value=float(st.session_state.get('ph96', 0.0))  # Ensure it's always a float
        )

        # Radio button for attending bootcamp via employer
        attending_via_employer = st.radio(
            "Are you attending this bootcamp via your current employer (has applicant been sent on bootcamp through their current employment)?",
            options=["Yes", "No"],
            index=0 if st.session_state.get('ph97y') == 'X' else 1  # Default based on previous selection
        )

        # Store selections
        if attending_via_employer == 'Yes':
            st.session_state.ph97y, st.session_state.ph97n = 'X', ''
        else:
            st.session_state.ph97y, st.session_state.ph97n = '', 'X'

        # Save additional inputs in session state
        st.session_state.ph93 = name_of_employer
        st.session_state.ph94 = employer_postcode
        st.session_state.ph95 = current_job_role
        st.session_state.ph96 = current_hourly_rate

        # Define radio button options for working alongside bootcamp
        work_alongside_bootcamp_options = {
            "Yes (Full-time employment)": "ph98",
            "Yes (Part-time employed)": "ph99",
            "Yes (Self-employed)": "ph100",
            "No": "ph101"
        }

        # Display radio buttons and store 'X' in the selected placeholder
        selected_work_plan = st.radio(
            "Do you plan to work alongside the bootcamp?",
            list(work_alongside_bootcamp_options.keys()),
            index=list(work_alongside_bootcamp_options.keys()).index(next(key for key, value in work_alongside_bootcamp_options.items() if st.session_state.get(value) == 'X')) if any(st.session_state.get(value) == 'X' for value in work_alongside_bootcamp_options.values()) else 0
        )

        # Update the corresponding placeholder with 'X'
        if selected_work_plan:
            # Reset all placeholders before setting the current one
            for key in work_alongside_bootcamp_options.values():
                st.session_state[key] = ''  # Reset all options to empty string
            st.session_state[work_alongside_bootcamp_options[selected_work_plan]] = 'X'  # Set selected option to 'X'

    if selected_employment in unpaid_employment_options:
        # Define radio button options for unemployment duration
        unemployment_options = {
            "Less than 6 months": "ph83",
            "6-11 months": "ph84",
            "12-23 months": "ph85",
            "24-35 months": "ph86",
            "36 months or over": "ph87"
        }

        # Display radio buttons and store 'X' in the selected placeholder
        selected_unemployment = st.radio(
            "If you are unemployed, how long have you been unemployed? (Please tick ONE box)",
            list(unemployment_options.keys()),
            index=list(unemployment_options.keys()).index(
                next((key for key, value in unemployment_options.items() if st.session_state.get(value) == 'X'), None)
            ) if any(st.session_state.get(value) == 'X' for value in unemployment_options.values()) else 0
        )

        # Update the corresponding placeholder with 'X'
        if selected_unemployment:
            # Reset all placeholders before setting the current one
            for key in unemployment_options.values():
                st.session_state[key] = ''  # Reset all options to empty string
            st.session_state[unemployment_options[selected_unemployment]] = 'X'  # Set selected option to 'X'

        # Define radio button options for benefits
        benefit_options = {
            "In receipt of JSA": "ph88",
            "In receipt of ESA (Part of WRAG group)": "ph89",
            "In receipt of Universal Credit": "ph90",
            "In receipt of another State Benefit": "ph91",
            "None": "ph92"
        }

        # Display radio buttons and store 'X' in the selected placeholder
        selected_benefit = st.radio(
            "If unemployed, please state what benefit you receive (Please tick ONE box)",
            list(benefit_options.keys()),
            index=list(benefit_options.keys()).index(
                next((key for key, value in benefit_options.items() if st.session_state.get(value) == 'X'), None)
            ) if any(st.session_state.get(value) == 'X' for value in benefit_options.values()) else 0
        )

        # Update the corresponding placeholder with 'X'
        if selected_benefit:
            # Reset all placeholders before setting the current one
            for key in benefit_options.values():
                st.session_state[key] = ''  # Reset all options to empty string
            st.session_state[benefit_options[selected_benefit]] = 'X'  # Set selected option to 'X'

        # Define radio button options for recent occupation
        recent_occupation_options = {
            "Major Group": "ph102",
            "Managers, directors and senior officials": "ph103",
            "Professional occupations": "ph104",
            "Associate professional and technical occupations": "ph105",
            "Administrative and secretarial occupations": "ph106",
            "Skilled trades occupations": "ph107",
            "Caring, leisure and other service occupations": "ph108",
            "Sales and customer service occupations": "ph109",
            "Process, plant and machine operatives": "ph110",
            "Elementary occupations": "ph111"
        }

        # Display radio buttons and store 'X' in the selected placeholder
        selected_recent_occupation = st.radio(
            "Please give your most recent occupation:",
            list(recent_occupation_options.keys()),
            index=list(recent_occupation_options.keys()).index(
                next((key for key, value in recent_occupation_options.items() if st.session_state.get(value) == 'X'), None)
            ) if any(st.session_state.get(value) == 'X' for value in recent_occupation_options.values()) else 0
        )

        # Update the corresponding placeholder with 'X'
        if selected_recent_occupation:
            # Reset all placeholders before setting the current one
            for key in recent_occupation_options.values():
                st.session_state[key] = ''  # Reset all options to empty string
            st.session_state[recent_occupation_options[selected_recent_occupation]] = 'X'  # Set selected option to 'X'

        # Define radio button options for industry/sector
        industry_sector_options = {
            "Agriculture / forestry / fishing": "ph112",
            "Distribution / hotels / restaurants": "ph113",
            "Public admin / education / health": "ph114",
            "Banking / finance": "ph115",
            "Energy / water": "ph116",
            "Transport / communication": "ph117",
            "Construction": "ph118",
            "Manufacturing": "ph119",
            "Other services": "ph120"
        }

        # Display radio buttons and store 'X' in the selected placeholder
        selected_industry_sector = st.radio(
            "Industry/sector of current occupation (please give most recent occupation):",
            list(industry_sector_options.keys()),
            index=list(industry_sector_options.keys()).index(
                next((key for key, value in industry_sector_options.items() if st.session_state.get(value) == 'X'), None)
            ) if any(st.session_state.get(value) == 'X' for value in industry_sector_options.values()) else 0
        )

        # Update the corresponding placeholder with 'X'
        if selected_industry_sector:
            # Reset all placeholders before setting the current one
            for key in industry_sector_options.values():
                st.session_state[key] = ''  # Reset all options to empty string
            st.session_state[industry_sector_options[selected_industry_sector]] = 'X'  # Set selected option to 'X'
            
            if st.session_state[industry_sector_options[selected_industry_sector]] == 'ph120':
                st.session_state.ph120a = st.text_input('Please specify below', value=st.session_state.get('ph120a', ''))

    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click with validation
    if next_clicked:
        if selected_employment in paid_employment_options:
            if (st.session_state.ph93 and
                st.session_state.ph94 and
                st.session_state.ph95):
                st.session_state.step = 7
                st.experimental_rerun()
            else:
                st.warning("Please fill in all fields before proceeding.")
        else:
            st.session_state.step = 7
            st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 5  # Go back to the previous step (Section 4)
        st.experimental_rerun()


elif st.session_state.step == 7:
    st.title("> 6: Disability, Learning Difficulty and/or Health Problem")

    # Initialize placeholders for conditions
    st.session_state.update({f'ph{i}': st.session_state.get(f'ph{i}', '') for i in range(121, 147)})
    st.session_state.impactful_condition = st.session_state.get('impactful_condition', '')  # Most impactful condition
    st.session_state.confidential_interview = st.session_state.get('confidential_interview', '')  # Confidential interview request
    
    # Radio button for initial question
    difficulty_options = ["No", "Yes", "Other"]
    
    # Retrieve the previously selected value
    selected_difficulty = next((key for key in difficulty_options if st.session_state.get(f'ph12{difficulty_options.index(key) + 1}') == 'X'), None)

    # Display the radio buttons and set default value
    selected_difficulty = st.radio(
        "Do you consider that you have a learning difficulty, disability, or health problem?",
        difficulty_options,
        index=difficulty_options.index(selected_difficulty) if selected_difficulty else 0
    )

    # Store 'X' in the corresponding placeholder based on the selected option
    st.session_state.ph121, st.session_state.ph122, st.session_state.ph123 = '', '', ''  # Reset all before setting
    if selected_difficulty == "No":
        st.session_state.ph121 = 'X'
    elif selected_difficulty == "Yes":
        st.session_state.ph122 = 'X'
    elif selected_difficulty == "Other":
        st.session_state.ph123 = 'X'
        other_specify = st.text_input("Please specify:", value=st.session_state.ph123a)
        st.session_state.ph123a = other_specify

    # If "Yes" is selected, display individual checkboxes for each condition
    if selected_difficulty == "Yes":
        st.subheader("Please select the specific difficulty, disability, or health problem:")

        # Checkbox options with corresponding placeholders
        if st.checkbox("Epilepsy", value=st.session_state.ph124 == 'X'):
            st.session_state.ph124 = 'X'
        else:
            st.session_state.ph124 = ''

        if st.checkbox("Hearing Impairment", value=st.session_state.ph125 == 'X'):
            st.session_state.ph125 = 'X'
        else:
            st.session_state.ph125 = ''

        if st.checkbox("Diagnosed mental health condition", value=st.session_state.ph126 == 'X'):
            st.session_state.ph126 = 'X'
        else:
            st.session_state.ph126 = ''

        if st.checkbox("Moderate Learning Difficulty", value=st.session_state.ph127 == 'X'):
            st.session_state.ph127 = 'X'
        else:
            st.session_state.ph127 = ''

        if st.checkbox("Physical Disability", value=st.session_state.ph128 == 'X'):
            st.session_state.ph128 = 'X'
        else:
            st.session_state.ph128 = ''

        if st.checkbox("Other Specific Learning Difficulty (e.g. Dyspraxia)", value=st.session_state.ph129 == 'X'):
            st.session_state.ph129 = 'X'
        else:
            st.session_state.ph129 = ''

        if st.checkbox("Profound/Complex Disabilities", value=st.session_state.ph130 == 'X'):
            st.session_state.ph130 = 'X'
        else:
            st.session_state.ph130 = ''

        if st.checkbox("Severe Learning Difficulty", value=st.session_state.ph131 == 'X'):
            st.session_state.ph131 = 'X'
        else:
            st.session_state.ph131 = ''

        if st.checkbox("Social, Emotional & Behavioural Difficulties", value=st.session_state.ph132 == 'X'):
            st.session_state.ph132 = 'X'
        else:
            st.session_state.ph132 = ''

        if st.checkbox("Speech, Language and Communication needs", value=st.session_state.ph133 == 'X'):
            st.session_state.ph133 = 'X'
        else:
            st.session_state.ph133 = ''

        if st.checkbox("Temporary Disability after Illness or accident", value=st.session_state.ph134 == 'X'):
            st.session_state.ph134 = 'X'
        else:
            st.session_state.ph134 = ''

        if st.checkbox("Visual Impairment-excluding glasses/contact lenses", value=st.session_state.ph135 == 'X'):
            st.session_state.ph135 = 'X'
        else:
            st.session_state.ph135 = ''

        if st.checkbox("Prefer not to say", value=st.session_state.ph136 == 'X'):
            st.session_state.ph136 = 'X'
        else:
            st.session_state.ph136 = ''

        if st.checkbox("Are you a wheelchair user?", value=st.session_state.ph137 == 'X'):
            st.session_state.ph137 = 'X'
        else:
            st.session_state.ph137 = ''

        if st.checkbox("Allergy", value=st.session_state.ph138 == 'X'):
            st.session_state.ph138 = 'X'
        else:
            st.session_state.ph138 = ''

        if st.checkbox("Asperger’s Syndrome", value=st.session_state.ph139 == 'X'):
            st.session_state.ph139 = 'X'
        else:
            st.session_state.ph139 = ''

        if st.checkbox("Asthma", value=st.session_state.ph140 == 'X'):
            st.session_state.ph140 = 'X'
        else:
            st.session_state.ph140 = ''

        if st.checkbox("Autism Spectrum Condition", value=st.session_state.ph141 == 'X'):
            st.session_state.ph141 = 'X'
        else:
            st.session_state.ph141 = ''

        if st.checkbox("Cystic Fibrosis", value=st.session_state.ph142 == 'X'):
            st.session_state.ph142 = 'X'
        else:
            st.session_state.ph142 = ''

        if st.checkbox("Diabetes", value=st.session_state.ph143 == 'X'):
            st.session_state.ph143 = 'X'
        else:
            st.session_state.ph143 = ''

        if st.checkbox("Disability Affecting Mobility", value=st.session_state.ph144 == 'X'):
            st.session_state.ph144 = 'X'
        else:
            st.session_state.ph144 = ''

        if st.checkbox("Dyscalculia", value=st.session_state.ph145 == 'X'):
            st.session_state.ph145 = 'X'
        else:
            st.session_state.ph145 = ''

        if st.checkbox("Dyslexia", value=st.session_state.ph146 == 'X'):
            st.session_state.ph146 = 'X'
        else:
            st.session_state.ph146 = ''

        # Additional text input for most impactful condition
        impactful_condition = st.text_input(
            "If you have ticked more than one of the above, please state which disability, learning difficulty, and/or health problem impacts most on your learning:",
            value=st.session_state.impactful_condition
        )
        
        # Checkbox for confidential interview request
        confidential_interview = st.checkbox(
            "If you have a support need and would benefit from a confidential interview, please tick this box",
            value=st.session_state.confidential_interview == 'X'
        )

        # Save the text input and checkbox in session state
        st.session_state.impactful_condition = impactful_condition
        st.session_state.confidential_interview = 'X' if confidential_interview else ''

    # Navigation buttons
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        if selected_difficulty == "Other":
            if st.session_state.ph123a:
                st.session_state.step = 8
                st.experimental_rerun()
            else:
                st.warning("Please fill in all fields before proceeding.")
        elif selected_difficulty == "Yes":
            if any(st.session_state.get(f'ph{i}') == 'X' for i in range(124, 147)):
                st.session_state.step = 8
                st.experimental_rerun()
            else:
                st.warning("Please fill in all fields before proceeding.")
        else:
            st.session_state.step = 8
            st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 6  # Go back to the previous step (Section 5)
        st.experimental_rerun()

elif st.session_state.step == 8:
    st.title("> 7: Contact and Marketing Information")

    # Initialize placeholders
    st.session_state.update({f'ph{i}': st.session_state.get(f'ph{i}', '') for i in range(147, 154)})
    st.session_state.other_source = st.session_state.get('other_source', '')  # Default for other source specification

    # Question: How did you hear about us?
    st.text("How did you hear about us?")

    # Radio button options
    options = [
        "Employer", 
        "Job Centre", 
        "Social Media", 
        "Local Press", 
        "Search Engine", 
        "Friends / Family", 
        "Other Source"
    ]

    # Retrieve the previously selected option
    selected_option = next((key for key in options if st.session_state.get(f'ph147' if key == "Employer" else 
                                                                      f'ph148' if key == "Job Centre" else 
                                                                      f'ph149' if key == "Social Media" else 
                                                                      f'ph150' if key == "Local Press" else 
                                                                      f'ph151' if key == "Search Engine" else 
                                                                      f'ph152' if key == "Friends / Family" else 
                                                                      f'ph153' if key == "Other Source" else None) == 'X'), None)

    # Display the radio buttons and set the default selection
    selected_option = st.radio("Select an option", options, index=options.index(selected_option) if selected_option else 0)

    # Set the corresponding placeholder based on selection
    # Reset all placeholders first
    for i in range(147, 154):
        st.session_state[f'ph{i}'] = ''  # Reset all options to empty string

    if selected_option == "Employer":
        st.session_state.ph147 = 'X'
    elif selected_option == "Job Centre":
        st.session_state.ph148 = 'X'
    elif selected_option == "Social Media":
        st.session_state.ph149 = 'X'
    elif selected_option == "Local Press":
        st.session_state.ph150 = 'X'
    elif selected_option == "Search Engine":
        st.session_state.ph151 = 'X'
    elif selected_option == "Friends / Family":
        st.session_state.ph152 = 'X'
    elif selected_option == "Other Source":
        st.session_state.ph153 = 'X'
        st.session_state.other_source = st.text_input("Please specify other source:", value=st.session_state.other_source)  # Set default value

    # Navigation buttons
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        if selected_option == "Other Source":
            if st.session_state.other_source:
                st.session_state.step = 9
                st.experimental_rerun()
            else:
                st.warning("Please fill in all fields before proceeding.")
        else:
            st.session_state.step = 9
            st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 7  # Go back to the previous step (Section 6)
        st.experimental_rerun()


elif st.session_state.step == 9:
    st.title("> 8: Learner Declaration and Commitment")
    
    # Initialize placeholders
    st.session_state.update({f'ph{i}': st.session_state.get(f'ph{i}', '') for i in range(154, 160)})

    # Agreement and Confirmation Section
    st.subheader("Agreement and Confirmation")

    privacy_notice = """
    I confirm that initial assessment and information advice and guidance concerning the course has been provided to me, this included information about the course, its entry requirements, its suitability and the support which is available to me.

    I agree that the information given on this agreement is true, correct and completed to the best of my knowledge and I understand that Prevista has the right to cancel my enrolment if it is found that I have provided false or inaccurate information. I agree that this information can be used to process my data for any purposes connected with my studies or my health and safety whilst on the premises. This also includes any other contractual requirements and in particular to the disclosure of all the data on this form or otherwise collected about me to the DfE for the purposes noted in the Privacy Notice in section 9.

    I also agree with the below points relating to my chosen programme:
    - Take appropriate responsibility for my own learning, development and progression
    - Attend and undertake training required to achieve the Skills Bootcamp identified in Programme Details in the ILP
    - Promptly inform the Employer and/or Prevista if any matters or issues arise, or might arise, that will, or may, affect my learning, development and progression
    - At all times behave in a safe and responsible manner and in accordance with the statutory requirements of health and safety law relating to my responsibilities from time to time

    If you wish to raise a complaint about how we have handled your personal data email to Prevista or any other issues, please contact us with full details of your issue. If you are not satisfied how your complaint has been dealt with, please be aware of Authority’s Whistleblowing and Complaints policies and processes.
    Whistleblowing involves entering a 'whistleblowing' webform on the 'Contact the Department for Education' page, which can be found below:
    Complaints Procedure - Department for Education - Gov.uk
    Contact the Department for Education - Gov.uk. Whistleblowing entries for Skills Bootcamps must be clearly marked as 'Skills Bootcamps' and will submitted via the DfE's whistleblowing submission process and will be escalated to the relevant policy team.

    Your information may also be shared with other third parties for the above purposes, but only where the law allows it and the sharing is in compliance with data protection legislation. You can agree to be contacted for other purposes by ticking any of the following boxes:
    """
    
    with st.container(height=600, border=True):
        st.write(privacy_notice)

    # Checkboxes for contact preferences
    st.session_state.ph154 = 'X' if st.checkbox("About courses or learning opportunities", value=st.session_state.ph154 == 'X') else ''
    st.session_state.ph155 = 'X' if st.checkbox("For surveys and research", value=st.session_state.ph155 == 'X') else ''
    st.session_state.ph156 = 'X' if st.checkbox("By post", value=st.session_state.ph156 == 'X') else ''
    st.session_state.ph157 = 'X' if st.checkbox("By phone", value=st.session_state.ph157 == 'X') else ''
    st.session_state.ph158 = 'X' if st.checkbox("By email", value=st.session_state.ph158 == 'X') else ''

    # Consent to filming
    st.session_state.ph159 = 'X' if st.checkbox("I consent to being filmed for course development, evaluation, and marketing purposes", key='checkbox159', value=st.session_state.ph159 == 'X') else ''
    
    # Display the state for debugging
    # st.write({
    #     'ph154': st.session_state.ph154,  # About courses or learning opportunities
    #     'ph155': st.session_state.ph155,  # For surveys and research
    #     'ph156': st.session_state.ph156,  # By post
    #     'ph157': st.session_state.ph157,  # By phone
    #     'ph158': st.session_state.ph158,  # By email
    #     'ph159': st.session_state.ph159   # Consent to being filmed
    # })    

    st.write("Participant Signature")
    # st.session_state.signature = st.text_input("Participant Signature")
    canvas_result = st_canvas(
        fill_color="rgba(255, 255, 255, 1)",  
        stroke_width=5,
        stroke_color="rgb(0, 0, 0)",  # Black stroke color
        background_color="white",  # White background color
        width=400,
        height=150,
        drawing_mode="freedraw",
        key="canvas",
    )
    st.session_state.signature = canvas_result.image_data

    # Set today's date automatically and display it
    st.session_state.date = date.today().strftime("%d-%m-%Y")
    st.write(f"Date: **{st.session_state.date}**")

    # Submit button
    submit_clicked = st.button("Submit")

###############################

    # Handle Submit button click
    if submit_clicked:
        if is_signature_drawn(st.session_state.signature) and st.session_state.date:
            time.sleep(1)
            st.write("**Thank you for completing the enrollment form!**")
            st.warning("Please wait for SNOWFLAKES!", icon="🚨")
            time.sleep(1)

            st.session_state.submission_done = True
            st.session_state.step = 10
            st.experimental_rerun()
        else:
            st.warning("Please provide your signature before submitting.")

#111111111111111111
    # Add a warning before the back button
    st.info("If you go back, you will have to re-sign the form.")

    # Navigation buttons
    back_clicked = st.button("Back", disabled=st.session_state.submission_done)

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 8  # Go back to the previous step
        st.experimental_rerun()
#11111111111111111

elif st.session_state.step == 10:
    st.info('Still Processing. . . .', icon="ℹ️")
    time.sleep(1)


# ####################################################################################################################################

with st.spinner('Wait for it...'):
    # Generate and save the document if form is submitted
    if st.session_state.submission_done:
        # FILL TEMPLATE:
        placeholder_values = {
            'ph7': st.session_state.title,
            'ph1': st.session_state.first_name,
            'ph2': st.session_state.sir_name,
            'ph8': st.session_state.preferred_name,
            'ph9': st.session_state.previous_name,
            'ph55': st.session_state.home_address,
            'ph4': st.session_state.postcode,
            'ph56': st.session_state.previous_postcode_country,
            'ph34': st.session_state.dob.strftime("%d-%m-%Y"),
            
            'ph57': st.session_state.current_age,
            'ph3': st.session_state.ni_number,
            'ph58': st.session_state.home_number,
            'ph6': st.session_state.mobile_number,
            'ph5': st.session_state.email,

            'ph219': st.session_state.ethnicity_vars['ethnicity_31'],
            'ph220': st.session_state.ethnicity_vars['ethnicity_32'],
            'ph221': st.session_state.ethnicity_vars['ethnicity_33'],
            'ph222': st.session_state.ethnicity_vars['ethnicity_34'],
            'ph223': st.session_state.ethnicity_vars['ethnicity_35'],
            'ph224': st.session_state.ethnicity_vars['ethnicity_36'],
            'ph225': st.session_state.ethnicity_vars['ethnicity_37'],
            'ph226': st.session_state.ethnicity_vars['ethnicity_38'],
            'ph227': st.session_state.ethnicity_vars['ethnicity_39'],
            'ph228': st.session_state.ethnicity_vars['ethnicity_40'],
            'ph229': st.session_state.ethnicity_vars['ethnicity_41'],
            'ph230': st.session_state.ethnicity_vars['ethnicity_42'],
            'ph231': st.session_state.ethnicity_vars['ethnicity_43'],
            'ph232': st.session_state.ethnicity_vars['ethnicity_44'],
            'ph233': st.session_state.ethnicity_vars['ethnicity_45'],
            'ph234': st.session_state.ethnicity_vars['ethnicity_46'],
            'ph235': st.session_state.ethnicity_vars['ethnicity_47'],
            'ph236': st.session_state.ethnicity_vars['ethnicity_48'],
            'ph59': st.session_state.ph59,
            'ph60': st.session_state.ph60,
            'ph61': st.session_state.ph61,
            'ph62': st.session_state.ph62,
            # 'ph10': st.session_state.highest_education,
            # 'ph11': st.session_state.institution_name,
            # 'ph12': st.session_state.year_of_completion,
            # 'ph13': st.session_state.completed_level_6_or_above,
            # 'ph14': st.session_state.subject,
            # 'ph15': st.session_state.date_applied,
            # 'ph16': st.session_state.placeholder_16,
            # 'ph17': st.session_state.placeholder_17,
            # 'ph18': st.session_state.placeholder_18,
            # 'ph19': st.session_state.placeholder_19,
            'ph40': st.session_state.emergency_contact_name,
            'ph41': st.session_state.emergency_contact_relationship,
            'ph42': st.session_state.emergency_contact_phone,
            'ph43': st.session_state.home_tel_no,

            'ph63': st.session_state.ph63,
            'ph64': st.session_state.ph64,
            'ph65': st.session_state.ph65,
            'ph66': st.session_state.ph66,
            'ph67': st.session_state.ph67,
            'ph68': st.session_state.ph68,
            'ph69': st.session_state.ph69,
            'ph70': st.session_state.ph70,
            'ph71': st.session_state.ph71,
            'ph72': st.session_state.ph72,
            'ph73': st.session_state.ph73,
            'ph74': st.session_state.ph74,        

            'ph75': st.session_state.ph75,
            'ph76': st.session_state.ph76,
            'ph77': st.session_state.ph77,
            'ph78': st.session_state.ph78,
            'ph79a': st.session_state.ph79a,
            'ph79b': st.session_state.ph79b,
            'ph79c': st.session_state.ph79c,
            'ph79d': st.session_state.ph79d,
            'ph80': st.session_state.ph80,
            'ph81': st.session_state.ph81,
            'ph82': st.session_state.ph82,

            'ph83': st.session_state.ph83,  # Less than 6 months
            'ph84': st.session_state.ph84,  # 6-11 months
            'ph85': st.session_state.ph85,  # 12-23 months
            'ph86': st.session_state.ph86,  # 24-35 months
            'ph87': st.session_state.ph87,   # 36 months or over

            'ph88': st.session_state.ph88,  # In receipt of JSA
            'ph89': st.session_state.ph89,  # In receipt of ESA (Part of WRAG group)
            'ph90': st.session_state.ph90,  # In receipt of Universal Credit
            'ph91': st.session_state.ph91,  # In receipt of another State Benefit
            'ph92': st.session_state.ph92,   # None

            'ph93': st.session_state.ph93,  # Name of Employer
            'ph94': st.session_state.ph94,  # Postcode
            'ph95': st.session_state.ph95,  # Current Job Role
            'ph96': st.session_state.ph96,  # Current Hourly Rate
            'ph97y': st.session_state.ph97y,   # Attending Bootcamp via Employer (Yes/No)
            'ph97n': st.session_state.ph97n,   # Attending Bootcamp via Employer (Yes/No)

            'ph98': st.session_state.ph98,  # Yes (Full-time employment)
            'ph99': st.session_state.ph99,  # Yes (Part-time employed)
            'ph100': st.session_state.ph100,  # Yes (Self-employed)
            'ph101': st.session_state.ph101,  # No

            'ph102': st.session_state.ph102,  # Major Group
            'ph103': st.session_state.ph103,  # Managers, directors and senior officials
            'ph104': st.session_state.ph104,  # Professional occupations
            'ph105': st.session_state.ph105,  # Associate professional and technical occupations
            'ph106': st.session_state.ph106,  # Administrative and secretarial occupations
            'ph107': st.session_state.ph107,  # Skilled trades occupations
            'ph108': st.session_state.ph108,  # Caring, leisure and other service occupations
            'ph109': st.session_state.ph109,  # Sales and customer service occupations
            'ph110': st.session_state.ph110,  # Process, plant and machine operatives
            'ph111': st.session_state.ph111,   # Elementary occupations

            'ph112': st.session_state.ph112,  # Agriculture / forestry / fishing
            'ph113': st.session_state.ph113,  # Distribution / hotels / restaurants
            'ph114': st.session_state.ph114,  # Public admin / education / health
            'ph115': st.session_state.ph115,  # Banking / finance
            'ph116': st.session_state.ph116,  # Energy / water
            'ph117': st.session_state.ph117,  # Transport / communication
            'ph118': st.session_state.ph118,  # Construction
            'ph119': st.session_state.ph119,  # Manufacturing
            'ph120': st.session_state.ph120,   # Other services (Please specify below)
            'ph120a': st.session_state.ph120a,   # Other services (Specifiy)

            'ph121': st.session_state.ph121,  # No
            'ph122': st.session_state.ph122,  # Yes
            'ph123': st.session_state.ph123,  # Other
            'ph123a': st.session_state.ph123a,  # Other Specify
            'ph124': st.session_state.ph124,  # Epilepsy
            'ph125': st.session_state.ph125,  # Hearing Impairment
            'ph126': st.session_state.ph126,  # Diagnosed mental health condition
            'ph127': st.session_state.ph127,  # Moderate Learning Difficulty
            'ph128': st.session_state.ph128,  # Physical Disability
            'ph129': st.session_state.ph129,  # Other Specific Learning Difficulty (e.g. Dyspraxia)
            'ph130': st.session_state.ph130,  # Profound/Complex Disabilities
            'ph131': st.session_state.ph131,  # Severe Learning Difficulty
            'ph132': st.session_state.ph132,  # Social, Emotional & Behavioural Difficulties
            'ph133': st.session_state.ph133,  # Speech, Language and Communication needs
            'ph134': st.session_state.ph134,  # Temporary Disability after Illness or accident
            'ph135': st.session_state.ph135,  # Visual Impairment-excluding glasses/contact lenses
            'ph136': st.session_state.ph136,  # Prefer not to say
            'ph137': st.session_state.ph137,  # Are you a wheelchair user? + Other specify text
            'ph138': st.session_state.ph138,  # Allergy
            'ph139': st.session_state.ph139,  # Asperger’s Syndrome
            'ph140': st.session_state.ph140,  # Asthma
            'ph141': st.session_state.ph141,  # Autism Spectrum Condition
            'ph142': st.session_state.ph142,  # Cystic Fibrosis
            'ph143': st.session_state.ph143,  # Diabetes
            'ph144': st.session_state.ph144,  # Disability Affecting Mobility
            'ph145': st.session_state.ph145,  # Dyscalculia
            'ph146': st.session_state.ph146,  # Dyslexia
            'impactful_condition': st.session_state.impactful_condition,  # Most impactful condition
            'confidential_interview': st.session_state.confidential_interview,  # Confidential interview checkbox

            'ph147': st.session_state.ph147,  # Employer
            'ph148': st.session_state.ph148,  # Job Centre
            'ph149': st.session_state.ph149,  # Social Media
            'ph150': st.session_state.ph150,  # Local Press
            'ph151': st.session_state.ph151,  # Search Engine
            'ph152': st.session_state.ph152,  # Friends / Family
            'ph153': st.session_state.ph153,  # Other Source
            'other_source': st.session_state.other_source,  # Other Source (specified)

            'ph154': st.session_state.ph154,  # About courses or learning opportunities
            'ph155': st.session_state.ph155,  # For surveys and research
            'ph156': st.session_state.ph156,  # By post
            'ph157': st.session_state.ph157,  # By phone
            'ph158': st.session_state.ph158,  # By email
            'ph159': st.session_state.ph159,   # Consent to being filmed

            'ph35m': st.session_state.ph35m,
            'ph35f': st.session_state.ph35f,
            'ph50': date.today().strftime("%d-%m-%Y"),
        }

        template_file = "ph_skills_bootcamp.docx"
        modified_file = f"SkillsBootcamp_Form_Submission_{st.session_state.first_name}_{st.session_state.sir_name}.docx"

        signature_path = 'signature_image.png'
        signature_image = PILImage.fromarray(
            st.session_state.signature.astype('uint8'), 'RGBA')
        signature_image.save(signature_path)

        replace_placeholders(template_file, modified_file, placeholder_values, signature_path)



    # Email
        # Sender email credentials
        # Credentials: Streamlit host st.secrets
        # sender_email = st.secrets["sender_email"]
        # sender_password = st.secrets["sender_password"]

        sender_email = get_secret("sender_email")
        sender_password = get_secret("sender_password")

        # Credentials: Local env
        # load_dotenv()                                     # uncomment import of this library!
        # sender_email = os.getenv('EMAIL')
        # sender_password = os.getenv('PASSWORD')
        team_email = [sender_email]
        # team_email = ['muhammadoa@prevista.co.uk']
        # receiver_email = sender_email
        # receiver_email = 'muhammadoa@prevista.co.uk'

        learner_email = [st.session_state.email]
        
        subject_team = f"Skills_Bootcamp: {st.session_state.selected_option} {st.session_state.hear_about}_{st.session_state.hother_source}_{st.session_state.first_name}_{st.session_state.sir_name} Submission Date: {date.today()}"
        body_team = "Prevista Skills Bootcamp Form submitted. Please find attached files."

        subject_learner = "Thank You for Your Interest in The Skills Bootcamp!"
        body_learner = f"""
        <html>
        <body>
            <p>Dear {st.session_state.first_name} {st.session_state.sir_name},</p>

            <p>Thank you for expressing your interest in Bootcamp at PREVISTA. We are excited to guide you through the next steps of the enrollment process.</p>

            <p><strong>What’s Next?</strong></p>
            <ol>
                <li><strong>Enrollment Communication:</strong> One of our representatives will be contacting you within the next few days to complete your enrollment. Please keep an eye out for our message to finalize your registration details.</li>
                <li><strong>Course Start Date:</strong> Once your enrollment is confirmed, we will send you the schedule for the course start date.</li>
                <li><strong>Orientation Session:</strong> You will be invited to an orientation session where you can learn more about the platform, meet your instructors, and connect with other learners.</li>
            </ol>

            <p>If you have any immediate questions, please feel free to reach out to us at PrevistaAdmissions@prevista.co.uk.</p>

            <p>We look forward to speaking with you soon and welcoming you to our learning community!</p>

            <p>Best regards,</p>
            <p>Student Admissions Team<br>
            PREVISTA<br>
            PREPARING YOU TODAY FOR OPPORTUNITIES OF TOMORROW</p>
        </body>
        </html>
        """

        # Local file path
        local_file_path = modified_file

        # Send email to team with attachments
        if st.session_state.files or local_file_path:
            send_email_with_attachments(sender_email, sender_password, team_email, subject_team, body_team, st.session_state.files, local_file_path)
        
        # Send thank you email to learner
        send_email_with_attachments(sender_email, sender_password, learner_email, subject_learner, body_learner)

        st.success("Processing Complete!")
        st.write("Someone will get in touch with you soon.")
        st.snow()
        last()




# streamlit run app.py
# Dev : https://linkedin.com/in/osamatech786